# coding:utf-8
from gevent import monkey
from retrying import retry

monkey.patch_all()
from gevent.pool import Pool
import execjs
import json
import re
import gevent
import requests
from queue import Queue
from docx.enum.text import WD_LINE_SPACING
from pandas import Series
import pandas as pd
from numpy import *
import time
import matplotlib.pyplot as plt
import numpy as np
import seaborn as sns
import scipy.stats as st
import xalpha as xa
from docx import Document
from docx.shared import Cm, Pt
from datetime import datetime

document = Document()
styles = document.styles
# 选取 style，并设置 style 中的段落格式
style = styles['Body Text']
para_format = style.paragraph_format
para_format.left_indent = Pt(0)
para_format.space_before = Pt(3)  # 固定值18磅
para_format.space_after = Pt(3)
para_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
Weight_input = 0
Weight_output = 0
ls = []
# 获取今天的估算数据
class Today_data:
    def __init__(self,quek,bianhao):
        self.quek=quek
        self.bianhao = bianhao
    # 提取数据的估算和基本信息
    @retry
    def get_1(self):
        self.code = str(self.bianhao)
        url = 'http://fundgz.1234567.com.cn/js/%s.js' % self.code+'?rt='
        print(url)
        result = requests.get(url)  # 发送请求
        data = json.loads(re.match(".*?({.*}).*", result.text, re.S).group(1))
        return data
    def information(self, data):
        self.quek.put('##############基金详情##############')
        self.quek.put('基金编码：%s %s' % (data['fundcode'], data['name']))
        self.quek.put('基金名称：%s' % data['name'])
        self.quek.put('单位净值：%s' % data['dwjz'])
        self.quek.put('净值日期：%s' % data['jzrq'])
        self.quek.put('估算值：%s' % data['gsz'])
        self.quek.put('估算增量：%s%%' % data['gszzl'])
        self.quek.put('估值时间：%s' % data['gztime'])
        # data1 = xa.universal.get_fund_peb(data['fundcode'], data['jzrq'], threhold=0.3)
        # self.quek.put('PB估值法数据：%s' % data1)#PB数据越小表示值得投资，PB=PE*ROM，PE和ROM比较大，表示泡沫现象严重
# 获取历史的数据，历史增长率，单位净值等
class history:
    def __init__(self,quek,code):
        self.quek=quek
        self.fscode = code

    # 提取网页（与下面fscode同时使用）
    # 提取增长率，可以分析数据的期望的方差（与上面fscode同时使用）
    def getWorth(self):
        def getUrl():
            head = 'http://fund.eastmoney.com/pingzhongdata/'
            tail = '.js?v=' + time.strftime("%Y%m%d%H%M%S", time.localtime())
            print(head + self.fscode + tail)
            return head + self.fscode + tail
        # 用requests获取到对应的文件
        content = requests.get(getUrl())
        # self.quek.put(getUrl(fscode))
        # 使用execjs获取到相应的数据
        jsContent = execjs.compile(content.text)
        name = jsContent.eval('fS_name')
        code = jsContent.eval('fS_code')
        # 单位净值走势
        netWorthTrend = jsContent.eval('Data_netWorthTrend')
        # 累计净值走势
        ACWorthTrend = jsContent.eval('Data_ACWorthTrend')

        netWorth = []
        ACWorth = []
        purworth = []
        # 提取出里面的净值
        for dayWorth in netWorthTrend[::-1]:
            netWorth.append(dayWorth['y'])
            purworth.append(dayWorth['equityReturn'])  # 增长率

        for dayACWorth in ACWorthTrend[::-1]:
            ACWorth.append(dayACWorth[1])
        # self.quek.put(name, code,purworth)
        return netWorth, purworth, ACWorth, name


class history3:
    def __init__(self,code):
        self.code = code

    @retry
    def history_3(self):
        zzyl = xa.fundinfo(self.code)
        a = zzyl.price
        a = zzyl.price.set_index('date')
        # print(a['netvalue'].values)
        # date  netvalue  comment  totvalue
        # print(a['netvalue'],a['totvalue'])
        # print(type(list(a['netvalue'].values)))
        # print(list(a['netvalue'].values))
        return a['netvalue'].values,a['totvalue'].values,self.code
# 绘制实际的增长率分布图
class Draw:
    def __init__(self,quek,list, bianhao, name, bins1=100, pro=0.95, rwidth=0.5):
        self.quek=quek
        self.list = list
        self.bianhao = bianhao
        self.name = name
        self.bins1 = bins1
        self.pro = pro
        self.rwidth = rwidth

    def draw(self):  # 增长率的概率控制
        plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签
        plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
        x = Series(self.list)
        x.plot(kind='hist', bins=self.bins1, density=True, label='直方图', color="steelblue", edgecolor="black",
               rwidth=self.rwidth)
        sns.distplot(x, rug=True, hist=False)
        loc1 = mean(x)
        self.quek.put('历史增长率均值为{}：'.format(loc1))
        scale1 = std(x)
        self.quek.put('历史增长率标准差为{}：'.format(scale1))
        X_max = st.norm.ppf(self.pro, loc=loc1, scale=scale1)  # 均值为10，标准差为0.5，分位值在0.025处对应的数字
        X_min = st.norm.ppf(1 - self.pro, loc=loc1, scale=scale1)
        a = self.bianhao + self.name
        plt.title(a)
        # plt.show()
        plt.close("all")
        self.quek.put(
            '最小{}%分位值的数据为：{}\n最大{}%分位值的数据为：{}'.format((1 - self.pro) * 100, X_min, (self.pro) * 100, X_max))
        return X_max, X_min
# 计算周期得出结果
class Timedata:#注意数据是从现在到以前排序
    def __init__(self,quek,netWorth, ACWorth, name, bianhao, probability, period, profit, gsz1):
        self.quek=quek
        self.netWorth = netWorth
        self.ACWorth = ACWorth
        self.name = name
        self.bianhao = bianhao
        self.probability = probability
        self.period = period
        self.profit = profit
        self.gsz1 = gsz1
    def combin1(self):  # 输入编号,盈利概率和计算周期,期待收益率
        # bianhao = str(self.bianhao)
        # netWorth, purworth, ACWorth, name = getWorth(bianhao)
        # self.quek.put('增长率为：{}'.format(purworth))
        # self.quek.put(netWorth)
        # gsz1 = float(combin2(bianhao)['gsz'])
        gsz1 = float(self.gsz1)
        # print(gsz1)
        # print(type(gsz1))
        # print('networth{}'.format(netWorth))
        list2 = self.netWorth[0:self.period:1]
        list3 = self.netWorth[0:self.period * 6:1]  # 取出半年的数据
        list4 = self.netWorth[0:self.period * 3:1]
        listx = list(range(self.period * 3))  # 取出三个月数据进行计算
        listy = list4
        listy.reverse()
        list5=self.netWorth[0:180:1]
        # self.quek.put(listy)
        # self.quek.put('最近{}天排序：{}'.format(period,sorted(list2)))
        # num=(((min(list2)+max(list2))/2)+min(list2))/2          #计算实际的最大值和最小值之间距离最小值1/4之间的距离
        nummin = (max(list2) - min(list2)) * (1 - self.probability) + min(list2)
        df = Series(list3)
        df0 = Series(list2)
        self.quek.put('{}天数据均值为：{}'.format(self.period, mean(df0)))
        mean_6mounth = mean(df)
        nummax = max(list4)-(max(list4) - min(list4)) * (1 - self.probability)
        self.quek.put('{}天数据的最低值{}%概率值对应数据为:{}'.format(self.period, self.probability * 100, nummin))
        # Slope1=gsz1-list2[0]
        # Slope2=list2[0]-list2[1]
        ma1=(max(list5)-min(list5))/min(list5)
        if gsz1 <= nummin and ma1>=0.045:
            self.quek.put('{}天数据升序排序为：'.format(self.period))
            # print(sorted(list2))
            self.quek.put(str(sorted(list2)))
            self.quek.put('{}天数据实际排序为：'.format(self.period))
            # print(list2)
            self.quek.put(str(list2))
            self.quek.put('{}天数据的最低值{}%概率值对应数据为:{}'.format(self.period, self.probability * 100, nummin))
            self.quek.put('------------------>前一天的净值为：{}'.format(list2[0]))
            self.quek.put('------------------->可以投资')
            buylist.append(self.bianhao)
            global Weight_input
            Weight_input = Weight_input + 5
        elif gsz1 >= nummax and ma1>=0.045:
            percent = (gsz1 - nummax) / nummax
            self.quek.put(
                '{}天数据的均值{}，基准以上{}%对应数据为:{}'.format(self.period * 6, mean_6mounth, self.profit * 100, nummax))
            self.quek.put(
                '------------------>今天估算的净值为：{}超出预期目标{}百分比为:{}%，建议卖出总资产:{}%'.format(gsz1, nummax, percent * 100,
                                                                                    percent * 100 + self.profit * 100))
            self.quek.put('{}天数据升序排序为：'.format(self.period))
            # print(sorted(list3))
            self.quek.put('{}天数据实际排序为：'.format(self.period))
            print(list4)
            self.quek.put(str(list4))
            self.quek.put('------------------->注意卖出')
            buylist.append(self.bianhao)
            # global Weight_output
            # Weight_output = Weight_output + 5
        else:
            self.quek.put('------------------>等待机会')
    def analysis_3days(self):
        _3days = self.netWorth[0:4:1]
        global Weight_output
        global Weight_input
        self.quek.put('最近4天变化情况：')
        if _3days[0] >= _3days[1] and _3days[1] >= _3days[2] and _3days[2] >= _3days[3]:
            self.quek.put('最近四天持续上涨')

            Weight_output = Weight_output + 2
        elif _3days[0] <= _3days[1] and _3days[1] <= _3days[2] and _3days[2] <= _3days[3]:
            self.quek.put('最近四天持续下跌')

            Weight_input = Weight_input + 2
        elif _3days[0] >= _3days[1] and _3days[1] >= _3days[2]:
            self.quek.put('最近三天持续上涨')

            Weight_output = Weight_output + 1
        elif _3days[0] <= _3days[1] and _3days[1] <= _3days[2]:
            self.quek.put('最近三天持续下跌')
            Weight_input = Weight_input + 1
        else:
            self.quek.put('最近处于波动状态')
    # # 分析数据的期望收益均值，方差，
    # def analysis_data(bianhao):
    #     # 提取网页（与下面fscode同时使用）
    #     # 提取增长率，可以分析数据的期望的方差（与上面fscode同时使用）
    #     netWorth, purworth, ACWorth, name = getWorth(bianhao)
    #     return purworth
# 分析最近四天的数据情况
class avedata:
    def __init__(self,quek,list1, min):
        self.quek=quek
        self.list1 = list1
        self.min = min
    def feature_data(self):  # 计划分析数据本身特性，确定合适波动变换区间，对波动性行业进行特征分析
        # 分析连续天数的数据，最小天数min和最大天数max
        # self.quek.put(self.list1)
        self.list1 = self.list1[::-1]
        # self.quek.put(self.list1)
        # self.quek.put(len(self.list1))
        loc_1 = []
        for j in range(self.min, self.min + 1):
            a = []
            scale_1 = []
            for i in range(0, len(self.list1) - j + 1):
                a1 = mean(self.list1[0 + i:j + i:1])
                a.append(a1)
            # self.quek.put('{:*^30}'.format('这是连续{}天的均值数据和列表'.format(j)))
            # self.quek.put(a)
            # self.quek.put('{:#^30}'.format('这里总共{}个数据'.format(str(len(a)))))
            # x = Series(a)
            # loc1 = std(x)
            # loc_1.append(loc1)
            # draw(a,'001717','实验')
            # scale_1.append(scale1)
            # draw_xy(X_max)
            # draw_xy(X_min)
        # self.quek.put(loc_1)
        # draw_xy(loc_1)
        # draw_xy(scale_1)
        return a
# 利用6,12,30天数据均值进行分析买卖点
class BIAS:
    def __init__(self,quek,name,code,list1, list2, list3,list4, block=90):
        self.quek=quek
        self.code = code
        self.name = name
        self.list1 = list1[-block:]
        self.list2 = list2[-block:]
        self.list3 = list3[-block:]
        self.list4 = list4[-block:]
        self.block = block
        # self.quek.put('{:*^30}'.format('这是连续{}天的均值数据和列表'.format(len(self.list1))))
        if list1[-1] < list2[-1] and list1[-1] < list3[-1] and list1[-1] < list4[-1]:
            self.quek.put("{:*^30}".format('已经跌破3类均线，可以投入'))
            if abs(list1[-1] - list2[-1]) >= 1.5 * abs(list2[-1] - list3[-1]):
                self.quek.put('已经跌破均线距离的整倍数，属于较大跌落')
        elif list1[-1] > list2[-1] and list1[-1] > list3[-1] and list1[-1] > list4[-1]:
            self.quek.put('已经超过均线-------------------------------------')
            if abs(list1[-1] - list2[-1]) >= 1.5 * abs(list2[-1] - list3[-1]):
                self.quek.put('已经上涨超过均线距离的整倍数，属于较大上涨')
        else:
            self.quek.put('再等等-------------------------------------')
    def bias_many(self):
        plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签
        plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
        plt.figure(figsize=(15, 7))
        x = np.arange(self.block)
        plt.title(self.code + self.name)
        plt.plot(x, self.list1)
        plt.plot(x, self.list2)
        plt.plot(x, self.list3)
        plt.plot(x, self.list4)
        plt.legend(['一天均值', '短期均值', '中期均值', '长期均值'], loc='upper left')
        plt.savefig('.\\picture\\' + self.code + '.png')
        plt.close("all")
        # plt.show()
class Industry_Code:
    def Industry(self):
        filepath = r'行业1.csv'
        df = pd.read_csv(filepath, dtype=str)
        print(df)
        a = df.set_index('行业')
        print(a)  # 显示所有的数据
        industry = input('请输入需要计算的行业：')
        # self.quek.put(a.loc[industry].dropna().loc['数据1'])#相应行业的代码去除了无效数据
        list1 = a.loc[industry].dropna().values
        print('行业主要的代表基金有：{}'.format(list1))  # 相应行业的代码去除了无效数据
        return list1

def worker(queq):
    count=0
    while not queq.empty():
        if count==1:
            task = queq.get()
            result='%s' % (task)
            print(result)
            document.add_heading('{}'.format(str(result)))
        else:
            task = queq.get()
            result='%s' % (task)
            print(result)
            document.add_paragraph('{}'.format(str(result)))
        count+=1
    print('Quitting time!')

def funcc(quek,listfullk1):
    code=listfullk1
    m1 = Today_data(quek,code)
    data = m1.get_1()
    name=data['name']
    m1.information(data)  # 获取今天的估值
    p1 = history3(code)  # 获取历史数据
    netWorth,  ACWorth, code = p1.history_3()
    netWorth=list(netWorth)[::-1]
    ACWorth=list(ACWorth)[::-1]
    # k1, k2 = Draw(purworth, '110011', name).draw()
    # print('shijiadasdad{}'.format(netWorth[::-1]))
    p3 = Timedata(quek,netWorth, ACWorth, name, code, 0.99, 7, 0.38, data['gsz'])
    p3.combin1()
    p3.analysis_3days()
    # pgs = history3(code).history_3()
    pgs1 = netWorth
    # pgs1.reverse()
    # self.quek.put(netWorth)
    pgs1.insert(0, float(data['gsz']))
    # print(pgs1)
    # self.quek.put(pgs)
    p6 = avedata(quek,pgs1, 1).feature_data()
    p7 = avedata(quek,pgs1, 7).feature_data()
    p8 = avedata(quek,pgs1, 14).feature_data()
    p9 = avedata(quek,pgs1, 45).feature_data()#45天均线
    p10 = BIAS(quek,name,code,p6, p7, p8, p9,120).bias_many()
    quek.put(code)
    print('{}分析结束'.format(code))
if __name__=='__main__':
    t = datetime.now().strftime('%Y%m%d')
    start_time=time.time()
    buylist=[]
    listfull1 = ['009180', '165520', '011148', '007301', '001630', '005038', '010770', '110030', '161725',
                 '110022', '161726', '000376', '000307']

    num=len(listfull1)
    lsk=[]
    for i in range(1,num+1):
        name = 'queue' + str(i)
        locals()[name] = Queue()
        lsk.append(locals()[name])
    pool = Pool(50)
    gevent.joinall([pool.spawn(funcc,lsk[i],listfull1[i]) for i in range(0,num)])
    for i in range(0,num):
        worker(lsk[i])      #注意这里的变量对象必须是使用locals()产生的
        document.add_picture('.\\picture\\' + listfull1[i] + '.png', width=Cm(15.24))
        document.add_page_break()
    document.save('.\\article\\'+t+'所有基金分析结果.docx')
    end_time=time.time()
    print(end_time-start_time)
    print(buylist)
    with open('待分析基金.txt','w') as f:
        for line in buylist:
            f.write(line+'\n')
