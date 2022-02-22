# *_*coding:utf-8 *_*
# *_*coding:utf-8 *_*
# *_*coding:utf-8 *_*
# coding:utf-8
import time
from collections import deque
from retrying import retry
from gevent import monkey

monkey.patch_all()
from gevent.pool import Pool
import execjs
import json
import re
import gevent
import requests
from queue import Queue
from docx.enum.text import WD_LINE_SPACING
from pandas import Series, DataFrame
import pandas as pd
from numpy import *
import time
import matplotlib.pyplot as plt
import numpy as np
import seaborn as sns
import scipy.stats as st
import xalpha as xa
import sys
from docx import Document
from docx.shared import Inches, Cm, Pt
from concurrent.futures import ThreadPoolExecutor, wait, ALL_COMPLETED
from openpyxl import Workbook

# 实例化
wb = Workbook()
# 激活 worksheet
ws = wb.active
from datetime import datetime

document = Document()
styles = document.styles
# 选取 style，并设置 style 中的段落格式
style = styles['Body Text']
para_format = style.paragraph_format
para_format.left_indent = Pt(0)
para_format.space_before = Pt(3)  # 固定值18磅
para_format.space_after = Pt(3)
para_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
Weight_input = 0
Weight_output = 0
ls = []


# 获取今天的估算数据
class Today_data:
    def __init__(self, quek, bianhao):
        self.quek = quek
        self.bianhao = bianhao

    # 提取数据的估算和基本信息
    @retry
    def get_1(self):
        self.code = str(self.bianhao)
        url = 'http://fundgz.1234567.com.cn/js/%s.js' % self.code + '?rt=1463558676006'
        print(url)
        result = requests.get(url)  # 发送请求
        data = json.loads(re.match(".*?({.*}).*", result.text, re.S).group(1))
        return data

    def information(self, data):
        self.quek.put('##############基金详情##############')
        self.quek.put('基金编码：%s %s' % (data['fundcode'], data['name']))
        self.quek.put('基金名称：%s' % data['name'])
        self.quek.put('单位净值：%s' % data['dwjz'])
        self.quek.put('净值日期：%s' % data['jzrq'])
        self.quek.put('估算值：%s' % data['gsz'])
        self.quek.put('估算增量：%s%%' % data['gszzl'])
        self.quek.put('估值时间：%s' % data['gztime'])
        # data1 = xa.universal.get_fund_peb(data['fundcode'], data['jzrq'], threhold=0.3)
        # self.quek.put('PB估值法数据：%s' % data1)#PB数据越小表示值得投资，PB=PE*ROM，PE和ROM比较大，表示泡沫现象严重


# 获取历史的数据，历史增长率，单位净值等
class history:
    def __init__(self, quek, code):
        self.quek = quek
        self.fscode = code

    # 提取网页（与下面fscode同时使用）
    # 提取增长率，可以分析数据的期望的方差（与上面fscode同时使用）
    def getWorth(self):
        def getUrl():
            head = 'http://fund.eastmoney.com/pingzhongdata/'
            tail = '.js?v=' + time.strftime("%Y%m%d%H%M%S", time.localtime())
            print(head + self.fscode + tail)
            return head + self.fscode + tail

        # 用requests获取到对应的文件
        content = requests.get(getUrl())
        # self.quek.put(getUrl(fscode))
        # 使用execjs获取到相应的数据
        jsContent = execjs.compile(content.text)
        name = jsContent.eval('fS_name')
        code = jsContent.eval('fS_code')
        # 单位净值走势
        netWorthTrend = jsContent.eval('Data_netWorthTrend')
        # 累计净值走势
        ACWorthTrend = jsContent.eval('Data_ACWorthTrend')

        netWorth = []
        ACWorth = []
        purworth = []
        # 提取出里面的净值
        for dayWorth in netWorthTrend[::-1]:
            netWorth.append(dayWorth['y'])
            purworth.append(dayWorth['equityReturn'])  # 增长率

        for dayACWorth in ACWorthTrend[::-1]:
            ACWorth.append(dayACWorth[1])
        # self.quek.put(name, code,purworth)
        return purworth


class history3:
    def __init__(self, code):
        self.code = code

    @retry
    def history_3(self):
        zzyl = xa.fundinfo(self.code)
        a = zzyl.price
        a = zzyl.price.set_index('date')
        # print(a['netvalue'].values)
        # date  netvalue  comment  totvalue
        # print(a['netvalue'],a['totvalue'])
        # print(type(list(a['netvalue'].values)))
        # print(list(a['netvalue'].values))
        return a['netvalue'].values, a['totvalue'].values, self.code


# 绘制实际的增长率分布图
class Draw:
    def __init__(self, data_today, quek, list, bianhao, name, bins1=100, pro=0.95, rwidth=0.5):
        self.data_today = data_today
        self.quek = quek
        self.list = list
        self.bianhao = bianhao
        self.name = name
        self.bins1 = bins1
        self.pro = pro
        self.rwidth = rwidth

    def draw(self):  # 增长率的概率控制
        fig1 = plt.figure()
        # plt.figure(figsize=(20, 25))
        plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签
        plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
        x = Series(self.list)
        x.plot(kind='hist', bins=self.bins1, density=True, label='直方图', color="steelblue", edgecolor="black",
               rwidth=self.rwidth)
        # sns.displot(x, rug=True)
        loc1 = mean(x)
        self.quek.put('历史增长率均值为{}：'.format(loc1))
        scale1 = std(x)
        self.quek.put('历史增长率标准差为{}：'.format(scale1))
        X_max = st.norm.ppf(self.pro, loc=loc1, scale=scale1)  # 均值为10，标准差为0.5，分位值在0.025处对应的数字
        X_min = st.norm.ppf(1 - self.pro, loc=loc1, scale=scale1)
        Percentage = st.norm.cdf(self.data_today, loc=loc1, scale=scale1)  # 标准正态分布在 0 处的累计分布概率
        a = self.bianhao + self.name
        plt.title(a)
        plt.savefig('.\\picture\\' + self.bianhao + '1.png')
        # plt.cla()`
        # plt.clf()
        # plt.close(fig1)

        self.quek.put(
            '最小{}%分位值的数据为：{}最大{}%分位值的数据为：{}'.format((1 - self.pro) * 100, X_min, (self.pro) * 100, X_max))
        self.quek.put(
            '当前增长率位于:{}'.format(Percentage))
        return X_max, X_min, Percentage


class Draw2:
    def __init__(self, data_today, quek, list, bianhao, name, bins1=100, pro=0.95, rwidth=0.5):
        self.data_today = data_today
        self.quek = quek
        self.list = list
        self.bianhao = bianhao
        self.name = name
        self.bins1 = bins1
        self.pro = pro
        self.rwidth = rwidth

    def draw2(self):  # 增长率的概率控制
        fig2 = plt.figure()
        plt.figure(figsize=(9, 15))
        plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签
        plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
        x = Series(self.list)
        x.plot(kind='hist', bins=self.bins1, density=True, label='直方图', color="steelblue", edgecolor="black",
               rwidth=self.rwidth)
        # sns.displot(x, rug=True)
        loc1 = mean(x)
        # self.quek.put('历史增长率均值为{}：'.format(loc1))
        scale1 = std(x)
        # self.quek.put('历史增长率标准差为{}：'.format(scale1))
        X_max = st.norm.ppf(self.pro, loc=loc1, scale=scale1)  # 均值为10，标准差为0.5，分位值在0.025处对应的数字
        X_min = st.norm.ppf(1 - self.pro, loc=loc1, scale=scale1)
        Percentage = st.norm.cdf(self.data_today, loc=loc1, scale=scale1)  # 标准正态分布在 0 处的累计分布概率
        a = self.bianhao + self.name
        plt.title(a)
        plt.savefig('.\\picture\\' + self.bianhao + '5.png')
        # plt.cla()
        plt.clf()
        plt.close(fig2)
        self.quek.put(
            '最小{}%分位值的数据为：{}最大{}%分位值的数据为：{}'.format((1 - self.pro) * 100, X_min, (self.pro) * 100, X_max))
        self.quek.put(
            '当前增长率位于:{}'.format(Percentage))
        return X_max, X_min, Percentage


# 计算周期得出结果
class Timedata:  # 注意数据是从现在到以前排序
    def __init__(self, quek, netWorth, ACWorth, name, bianhao, probability, period, profit, gsz1):
        self.quek = quek
        self.netWorth = netWorth
        self.ACWorth = ACWorth
        self.name = name
        self.bianhao = bianhao
        self.probability = probability
        self.period = period
        self.profit = profit
        self.gsz1 = gsz1

    def combin1(self):  # 输入编号,盈利概率和计算周期,期待收益率
        # bianhao = str(self.bianhao)
        # netWorth, purworth, ACWorth, name = getWorth(bianhao)
        # self.quek.put('增长率为：{}'.format(purworth))
        # self.quek.put(netWorth)
        # gsz1 = float(combin2(bianhao)['gsz'])
        gsz1 = float(self.gsz1)
        # print(gsz1)
        # print(type(gsz1))
        # print('networth{}'.format(netWorth))
        list2 = self.netWorth[0:self.period:1]
        list3 = self.netWorth[0:self.period * 6:1]  # 取出半年的数据
        list4 = self.netWorth[0:self.period * 3:1]
        listx = list(range(self.period * 3))  # 取出三个月数据进行计算
        listy = list4
        listy.reverse()
        list5 = self.netWorth[0:180:1]
        # self.quek.put(listy)
        # self.quek.put('最近{}天排序：{}'.format(period,sorted(list2)))
        # num=(((min(list2)+max(list2))/2)+min(list2))/2          #计算实际的最大值和最小值之间距离最小值1/4之间的距离
        nummin = (max(list5) - min(list5)) * (1 - self.probability) + min(list2)
        df = Series(list3)
        df0 = Series(list2)
        self.quek.put('{}天数据均值为：{}'.format(self.period, mean(df0)))
        mean_6mounth = mean(df0)
        nummax = max(list5) - (max(list5) - min(list5)) * (1 - self.probability)
        self.quek.put('{}天数据的最低值{}%概率值对应数据为:{}'.format(self.period, self.probability * 100, nummin))
        # Slope1=gsz1-list2[0]
        # Slope2=list2[0]-list2[1]
        ma1 = abs((max(list5) - min(list5)) / min(list5))
        percent = (gsz1 - nummin) / (nummax - nummin)
        if percent >= 1.4:
            word1 = '卖出100%'
        elif percent >= 1.3:
            word1 = '卖出45%'
        elif percent >= 1.2:
            word1 = '卖出23%'
        elif percent >= 1.1:
            word1 = '卖出13%'
        elif percent >= 1.0:
            word1 = '卖出9%'
        elif percent >= 0:
            word1 = '当前位置位于历史最高和最低点区间{}%'.format(percent*100)
        if gsz1 <= nummin and ma1 >= 0.07:
            self.quek.put('{}天数据升序排序为：{}'.format(self.period, str(sorted(list2))))
            # print(sorted(list2))
            self.quek.put('{}天数据实际排序为：{}'.format(self.period, str(list2)))
            # print(list2)
            self.quek.put('{}天数据的最低值{}%概率值对应数据为:{}'.format(self.period, self.probability * 100, nummin))
            self.quek.put('------------------>前一天的净值为：{}'.format(list2[0]))
            self.quek.put('------------------->可以投资')
            global Weight_input
            Weight_input = Weight_input + 5
        elif gsz1 >= nummax and ma1 >= 0.07:
            self.quek.put(
                '{}天数据的均值{}，基准以上{}%对应数据为:{}'.format(self.period * 3, mean_6mounth, self.profit * 100, nummax))
            self.quek.put(
                '------------------>今天估算的净值为：{}超出最高预期目标{}百分比为:{}%，建议卖出:{}'.format(gsz1, nummax, percent * 100,word1))
            self.quek.put('{}天数据升序排序为：{}'.format(self.period, str(sorted(list4))))
            self.quek.put('{}天数据实际排序为：{}'.format(self.period, str(list4)))
            self.quek.put('------------------->注意卖出')
        else:
            self.quek.put(' ')
            self.quek.put(' ')
            self.quek.put(' ')
            self.quek.put(word1)
            self.quek.put('------------------>等待机会')

    def analysis_3days(self):
        _3days = self.netWorth[0:4:1]
        global Weight_output
        global Weight_input
        self.quek.put('最近4天变化情况：')
        if _3days[0] >= _3days[1] and _3days[1] >= _3days[2] and _3days[2] >= _3days[3]:
            self.quek.put('最近四天持续上涨')
            Weight_output = Weight_output + 2
        elif _3days[0] <= _3days[1] and _3days[1] <= _3days[2] and _3days[2] <= _3days[3]:
            self.quek.put('最近四天持续下跌')
            Weight_input = Weight_input + 2
        elif _3days[0] >= _3days[1] and _3days[1] >= _3days[2]:
            self.quek.put('最近三天持续上涨')
            Weight_output = Weight_output + 1
        elif _3days[0] <= _3days[1] and _3days[1] <= _3days[2]:
            self.quek.put('最近三天持续下跌')
            Weight_input = Weight_input + 1
        else:
            self.quek.put('最近处于波动状态')
    # # 分析数据的期望收益均值，方差，
    # def analysis_data(bianhao):
    #     # 提取网页（与下面fscode同时使用）
    #     # 提取增长率，可以分析数据的期望的方差（与上面fscode同时使用）
    #     netWorth, purworth, ACWorth, name = getWorth(bianhao)
    #     return purworth


# 分析最近四天的数据情况
class avedata:
    def __init__(self, quek, list1, min):
        self.quek = quek
        self.list1 = list1
        self.min = min

    def feature_data(self):  # 计划分析数据本身特性，确定合适波动变换区间，对波动性行业进行特征分析
        # 分析连续天数的数据，最小天数min和最大天数max
        # self.quek.put(self.list1)
        self.list1 = self.list1[::-1]
        # self.quek.put(self.list1)
        # self.quek.put(len(self.list1))
        loc_1 = []
        for j in range(self.min, self.min + 1):
            a = []
            scale_1 = []
            for i in range(0, len(self.list1) - j + 1):
                a1 = mean(self.list1[0 + i:j + i:1])
                a.append(a1)
            # self.quek.put('{:*^30}'.format('这是连续{}天的均值数据和列表'.format(j)))
            # self.quek.put(a)
            # self.quek.put('{:#^30}'.format('这里总共{}个数据'.format(str(len(a)))))
            # x = Series(a)
            # loc1 = std(x)
            # loc_1.append(loc1)
            # draw(a,'001717','实验')
            # scale_1.append(scale1)
            # draw_xy(X_max)
            # draw_xy(X_min)
        self.quek.put('平均{}天增长幅度为{}'.format(self.min, a[-1]))
        # draw_xy(loc_1)
        # draw_xy(scale_1)
        return a

    def feature_data1(self):  # 计划分析数据本身特性，确定合适波动变换区间，对波动性行业进行特征分析
        # 分析连续天数的数据，最小天数min和最大天数max
        # self.quek.put(self.list1)
        self.list1 = self.list1[::-1]
        # self.quek.put(self.list1)
        # self.quek.put(len(self.list1))
        loc_1 = []
        for j in range(self.min, self.min + 1):
            a = []
            scale_1 = []
            for i in range(0, len(self.list1) - j + 1):
                a1 = sum(self.list1[0 + i:j + i:1])
                a.append(a1)
            # self.quek.put('{:*^30}'.format('这是连续{}天的均值数据和列表'.format(j)))
            # self.quek.put(a)
            # self.quek.put('{:#^30}'.format('这里总共{}个数据'.format(str(len(a)))))
            # x = Series(a)
            # loc1 = std(x)
            # loc_1.append(loc1)
            # draw(a,'001717','实验')
            # scale_1.append(scale1)
            # draw_xy(X_max)
            # draw_xy(X_min)
        self.quek.put('累计{}天增长幅度为{}'.format(self.min, a[-1]))
        # draw_xy(loc_1)
        # draw_xy(scale_1)
        return a

    def feature_data2(self):  # 计划分析数据本身特性，确定合适波动变换区间，对波动性行业进行特征分析
        # 分析连续天数的数据，最小天数min和最大天数max
        # self.quek.put(self.list1)
        self.list1 = self.list1[::-1]
        # self.quek.put(self.list1)
        # self.quek.put(len(self.list1))
        loc_1 = []
        for j in range(self.min, self.min + 1):
            a = []
            scale_1 = []
            for i in range(0, len(self.list1) - j + 1):
                a1 = mean(self.list1[0 + i:j + i:1])
                a.append(a1)
            # self.quek.put('{:*^30}'.format('这是连续{}天的均值数据和列表'.format(j)))
            # self.quek.put(a)
            # self.quek.put('{:#^30}'.format('这里总共{}个数据'.format(str(len(a)))))
            # x = Series(a)
            # loc1 = std(x)
            # loc_1.append(loc1)
            # draw(a,'001717','实验')
            # scale_1.append(scale1)
            # draw_xy(X_max)
            # draw_xy(X_min)
        self.quek.put('{}天均值数据为{}'.format(self.min, a[-1]))
        # draw_xy(loc_1)
        # draw_xy(scale_1)
        return a


# 利用6,12,30天数据均值进行分析买卖点
class BIAS:
    def __init__(self, i, quek, name, code, list1, list2, list3, list4, block=90):
        self.i = i
        self.quek = quek
        self.code = code
        self.name = name
        self.list1 = list1[-block:]
        self.list2 = list2[-block:]
        self.list3 = list3[-block:]
        self.list4 = list4[-block:]
        self.block = block
        # self.quek.put('{:*^30}'.format('这是连续{}天的均值数据和列表'.format(len(self.list1))))
        # if list1[-1] < list2[-1] and list1[-1] < list3[-1] and list1[-1] < list4[-1]:
        #     self.quek.put("{:*^30}".format('已经跌破3类均线，可以投入'))
        #     if abs(list1[-1] - list2[-1]) >= 1.5 * abs(list2[-1] - list3[-1]):
        #         self.quek.put('已经跌破均线距离的整倍数，属于较大跌落')
        # elif list1[-1] > list2[-1] and list1[-1] > list3[-1] and list1[-1] > list4[-1]:
        #     self.quek.put('已经超过均线-------------------------------------')
        #     if abs(list1[-1] - list2[-1]) >= 1.5 * abs(list2[-1] - list3[-1]):
        #         self.quek.put('已经上涨超过均线距离的整倍数，属于较大上涨')
        # else:
        #     self.quek.put('再等等-------------------------------------')

    def bias_many(self):
        fig3 = plt.figure()
        plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签
        plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
        plt.figure(figsize=(15, 9))
        x = np.arange(self.block)
        plt.title(self.code + self.name)
        plt.plot(x, self.list1)
        plt.plot(x, self.list2)
        plt.plot(x, self.list3)
        plt.plot(x, self.list4)
        plt.legend(['一天均值', '短期均值', '中期均值', '长期均值'], loc='upper left')
        plt.savefig('.\\picture\\' + self.code + str(self.i) + '.png')
        plt.clf()
        plt.close('all')
        # plt.show()


def worker(queq):
    count = 0
    ls = []
    while not queq.empty():
        if count == 1:
            task = queq.get()
            result = '%s' % (task)
            print(result)
            ls.append('{}'.format(str(result)))
            document.add_heading('{}'.format(str(result)))
        else:
            task = queq.get()
            result = '%s' % (task)
            print(result)
            if count not in [0, 2, 3, 4, 5, 13, 14, 15, 16, 17]:
                ls.append('{}'.format(str(result)))
            # ls.append('{}'.format(str(result)))
            document.add_paragraph('{}'.format(str(result)))
        count += 1
    ws.append(ls)
    print('Quitting time!')


def funcc(quek, listfullk1):
    code = listfullk1
    m1 = Today_data(quek, code)
    data = m1.get_1()
    name = data['name']
    m1.information(data)  # 获取今天的估值
    p1 = history3(code)  # 获取历史数据

    netWorth, ACWorth, code = p1.history_3()
    netWorth = list(netWorth)[::-1]
    print('-----',netWorth)
    ACWorth = list(ACWorth)[::-1]
    purworth = history(quek, code).getWorth()
    x_max, x_min, percentage = Draw(float(data['gszzl']), quek, purworth, code, name).draw()
    if float(data['gszzl']) <= float(x_min):
        quek.put('低点小概率事件发生，可以补仓')
    elif float(data['gszzl']) >= float(x_max):
        quek.put('高点小概率事件发生，注意撤仓')
    else:
        quek.put('数据区间正常')
    # print('实际增长率{}'.format(purworth[::-1]))
    p3 = Timedata(quek, netWorth, ACWorth, name, code, 0.98, 30, 0.38, data['gsz'])
    p3.combin1()
    p3.analysis_3days()
    # pgs = history3(code).history_3()
    pgs2 = purworth
    # pgs1.reverse()
    # self.quek.put(netWorth)
    pgs2.insert(0, float(data['gszzl']))
    # print(pgs1)
    # self.quek.put(pgs)
    p6 = avedata(quek, pgs2, 1).feature_data()
    p7 = avedata(quek, pgs2, 2).feature_data()
    p8 = avedata(quek, pgs2, 3).feature_data()
    p9 = avedata(quek, pgs2, 5).feature_data()  # 45天均线
    p10 = BIAS(2, quek, name, code, p6, p7, p8, p9, 180).bias_many()
    p6_2 = avedata(quek, pgs2, 2).feature_data1()
    p7_2 = avedata(quek, pgs2, 3).feature_data1()
    p8_2 = avedata(quek, pgs2, 5).feature_data1()
    p9_2 = avedata(quek, pgs2, 7).feature_data1()  # 45天总和线
    x_max_sum, x_min_sum, percentage_sum = Draw2(p9_2[-1], quek, p9_2, code, name).draw2()
    p10_2 = BIAS(3, quek, name, code, p6_2, p7_2, p8_2, p9_2, 180).bias_many()

    pgs1 = netWorth
    # pgs1.reverse()
    # self.quek.put(netWorth)
    pgs1.insert(0, float(data['gsz']))
    # print(pgs1)
    # self.quek.put(pgs)
    p6_1 = avedata(quek, pgs1, 1).feature_data2()
    p7_1 = avedata(quek, pgs1, 7).feature_data2()
    p8_1 = avedata(quek, pgs1, 14).feature_data2()
    p9_1 = avedata(quek, pgs1, 45).feature_data2()  # 45天均线
    p10_1 = BIAS(4, quek, name, code, p6_1, p7_1, p8_1, p9_1, 180).bias_many()
    quek.put(code)
    print('{}分析结束'.format(code))


if __name__ == '__main__':
    t = datetime.now().strftime('%Y%m%d')
    listfull1 = []
    start_time = time.time()
    with open('待分析基金.txt', 'r') as f:
        for line in f:
            listfull1.append(line.strip('\n'))
    print(listfull1)
    num = len(listfull1)
    lsk = []
    for i in range(1, num + 1):
        name = 'queue' + str(i)
        locals()[name] = Queue()
        lsk.append(locals()[name])
    pool = Pool(50)
    gevent.joinall([pool.spawn(funcc, lsk[i], listfull1[i]) for i in range(0, num)])
    for i in range(0, num):
        worker(lsk[i])  # 注意这里的变量对象必须是使用locals()产生的
        document.add_picture('.\\picture\\' + listfull1[i] + '1.png', width=Cm(15.24))
        document.add_picture('.\\picture\\' + listfull1[i] + '2.png', width=Cm(15.24))
        document.add_picture('.\\picture\\' + listfull1[i] + '3.png', width=Cm(15.24))
        document.add_picture('.\\picture\\' + listfull1[i] + '4.png', width=Cm(15.24))
        document.add_picture('.\\picture\\' + listfull1[i] + '5.png', width=Cm(15.24))
        document.add_page_break()
    document.save('.\\article\\' + t + '异常波动基金分析结果.docx')
    wb.save('.\\article\\' + t + '异常波动基金分析结果.xlsx')
    end_time = time.time()
    print(end_time - start_time)








